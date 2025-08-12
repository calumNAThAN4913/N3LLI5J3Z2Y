# 代码生成时间: 2025-08-13 00:24:19
import dash
import dash_core_components as dcc
import dash_html_components as html
from dash.dependencies import Input, Output
import plotly.express as px
from docx import Document
from io import BytesIO
import base64
import pandas as pd
import os

# Define the layout of the Dash app
app = dash.Dash(__name__)

app.layout = html.Div(children=[
    html.H1("Document Format Converter"),
    dcc.Upload(
        id="upload-data",
        children=html.Button("Upload Document"),
        style={"width": "100%", "height": "60px", "lineHeight": "60px"},
        # Allow multiple files to be uploaded
        multiple=True
    ),
    html.Div(id="output-data-upload"),
    dcc.Download(id="download-converted")
])

# Define a callback to update the output div when a file is uploaded
@app.callback(
    Output("output-data-upload", "children"),
    [Input("upload-data", "contents")],
    [State("upload-data", "filename")]
)
def update_output(contents, filename):
    if contents is None:
        return None
    try:
        # Open the document and read its content
        doc = Document(BytesIO(contents))
        # Convert the document to a pandas DataFrame
        df = pd.DataFrame([{
            "filename": filename,
            "text": paragraph.text
        } for paragraph in doc.paragraphs])
        # Display the first few rows of the DataFrame
        return html.Div(
            style={'whiteSpace': 'pre', 'fontFamily': 'monospace'},
            children=[
                html.P(f"Filename: {filename}"),
                html.Pre(df.head().to_string())
            ]
        )
    except Exception as e:
        return f"Error: {e}"

# Define a callback to download the converted document
@app.callback(
    Output("download-converted", "data"),
    [Input("upload-data", "contents")],
    [State("upload-data", "filename"), State("output-data-upload", "children")]
)
def download_file(contents, filename, children):
    if contents is None:
        return None
    try:
        # Get the content of the uploaded document
        doc = Document(BytesIO(contents))
        # Create a new document with the converted content
        new_doc = Document()
        for paragraph in doc.paragraphs:
            new_doc.add_paragraph(paragraph.text)
        # Save the new document to a BytesIO object
        output = BytesIO()
        new_doc.save(output)
        output.seek(0)
        # Return the converted document as a downloadable file
        return base64.b64encode(output.getvalue()).decode('utf-8')
    except Exception as e:
        return f"Error: {e}"

# Run the app
if __name__ == '__main__':
    app.run_server(debug=True)